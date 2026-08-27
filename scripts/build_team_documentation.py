from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "TEAM_FEATURE_DOCUMENTATION.md"
OUTPUT = ROOT / "docs" / "CampusHub_Team_Feature_Documentation.docx"
DIAGRAMS = ROOT / "docs" / "diagrams"

INK = "17212B"
CORAL = "BD4F3C"
SUN = "F4B83F"
PAPER = "F8F1E5"
SAGE = "DCE8DC"
MUTED = "5F696F"
WHITE = "FFFFFF"
GRID = "D7D1C6"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def create_numbering_instance(doc):
    root = doc.part.numbering_part.element
    abstract_id = None
    for abstract in root.findall(qn("w:abstractNum")):
        fmt = abstract.find(".//" + qn("w:numFmt"))
        if fmt is not None and fmt.get(qn("w:val")) == "decimal":
            abstract_id = abstract.get(qn("w:abstractNumId"))
            break
    if abstract_id is None:
        abstract_id = "0"
    ids = [int(n.get(qn("w:numId"))) for n in root.findall(qn("w:num"))]
    num_id = max(ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    root.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_rich_text(paragraph, text, size=None, color=None):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=(size or 10) - 0.5, color=CORAL)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def add_callout(doc, text, label="IMPORTANT"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9120])
    cell = table.cell(0, 0)
    set_repeat_table_header(table.rows[0])
    set_cell_shading(cell, SAGE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=8.5, color=CORAL, bold=True)
    add_rich_text(p, text, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, filename, caption, alt, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(DIAGRAMS / filename), width=Inches(width))
    drawing = run._element.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", alt)
        drawing[0].set("title", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(9)
    cr = cp.add_run(caption)
    set_font(cr, size=8.5, color=MUTED, italic=True)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    settings = {
        "Title": (30, INK, 0, 8),
        "Subtitle": (12.5, MUTED, 0, 14),
        "Heading 1": (19, INK, 15, 7),
        "Heading 2": (13.5, CORAL, 11, 4),
        "Heading 3": (10.8, INK, 8, 3),
    }
    for style_name, (size, color, before, after) in settings.items():
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Arial"
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.08


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("CAMPUSHUB  /  TEAM FEATURE DOCUMENTATION")
    set_font(r, size=7.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    left = fp.add_run("Faisal Mahbub  ·  Rifat Mahmud  ·  Tarannum Diha")
    set_font(left, size=7.5, color=MUTED)
    fp.add_run("\t")
    add_page_number(fp)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("CAMPUS CLUB & EVENT HUB")
    set_font(r, size=10, color=CORAL, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Team Feature\nDocumentation")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Architecture, ownership, demonstrations, and viva preparation")
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("━━━━━━━━━━━━━━━━━━━━")
    set_font(rr, size=12, color=SUN)
    owners = doc.add_paragraph()
    owners.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owners.paragraph_format.space_before = Pt(24)
    add_rich_text(owners, "Faisal Mahbub  ·  Rifat Mahmud  ·  Tarannum Diha", size=11, color=INK)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(70)
    add_rich_text(meta, "PHP 8  /  MySQL  /  XAMPP  /  August 2026", size=9, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    p = doc.add_paragraph("Document map", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    entries = [
        ("01", "Project and shared architecture"),
        ("02", "Database model and relationships"),
        ("03", "Tarannum Diha — management lifecycle"),
        ("04", "Rifat Mahmud — registration lifecycle"),
        ("05", "Faisal Mahbub — dashboards and engagement"),
        ("06", "Integration, status, demo, and viva reference"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for number, label in entries:
        cells = table.add_row().cells
        cells[0].text = number
        cells[1].text = label
        set_cell_shading(cells[0], SUN)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, size=9, color=INK, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, size=10, color=INK, bold=True)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, [800, 8320])
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()
    add_callout(
        doc,
        "All modules documented as implemented have working interfaces. QR, payments, external messaging, biometric attendance, and external verification services are intentionally outside scope.",
        "PRESENTATION RULE",
    )
    doc.add_page_break()


def clean_md(text):
    return text.replace("→", "->").replace("–", "-").replace("—", "-").replace("..", "..")


def add_markdown_table(doc, rows):
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) > 1 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in parsed[1]):
        parsed.pop(1)
    cols = len(parsed[0])
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, values in enumerate(parsed):
        cells = table.add_row().cells
        for ci, value in enumerate(values):
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, clean_md(value), size=8.2 if cols >= 4 else 8.8, color=WHITE if ri == 0 else INK)
            if ri == 0:
                set_cell_shading(cells[ci], INK)
                for run in p.runs:
                    run.bold = True
            elif ri % 2 == 0:
                set_cell_shading(cells[ci], PAPER)
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    widths = {
        2: [2800, 6320],
        3: [2100, 2250, 4770],
        4: [2150, 1450, 2600, 2920],
        5: [1700, 1350, 1600, 2150, 2320],
    }.get(cols, [9120 // cols] * cols)
    widths[-1] += 9120 - sum(widths)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_contents(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    # The DOCX has its own cover and front matter.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1. Project overview"))
    lines = lines[start:]
    i = 0
    in_code = False
    code_lines = []
    skip_code_block = False
    active_num_id = None
    figure_added = set()
    while i < len(lines):
        raw = lines[i]
        line = clean_md(raw.strip())
        if raw.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
                skip_code_block = False
            else:
                if not skip_code_block:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.22)
                    p.paragraph_format.right_indent = Inches(0.16)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(7)
                    r = p.add_run("\n".join(code_lines))
                    set_font(r, name="Consolas", size=8.1, color=INK)
                in_code = False
            i += 1
            continue
        if in_code:
            if not code_lines and raw.lstrip().startswith("USERS 1"):
                skip_code_block = True
            code_lines.append(clean_md(raw))
            i += 1
            continue
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            active_num_id = None
            level = len(heading_match.group(1))
            title = heading_match.group(2).replace("#", "").strip()
            # Create clear owner section openings.
            if title.startswith(("4. Tarannum", "5. Rifat", "6. Faisal", "7. Cross-feature")):
                doc.add_page_break()
            style = "Heading 1" if level <= 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_rich_text(p, title, color=INK if style != "Heading 2" else CORAL)
            if title.startswith("2.1 Request flow") and "architecture" not in figure_added:
                add_figure(doc, "system-architecture.png", "Figure 1. CampusHub request architecture", "Browser requests pass through Apache and PHP to the normalized MySQL database.")
                figure_added.add("architecture")
            elif title.startswith("3.1 Core relational chain") and "schema" not in figure_added:
                add_figure(doc, "normalized-schema.png", "Figure 2. Normalized core relational schema", "Core database entities with primary keys, foreign keys and relationship cardinalities.", 6.45)
                figure_added.add("schema")
            elif title.startswith("4.4 Main PHP") and "secure" not in figure_added:
                add_figure(doc, "secure-action-flow.png", "Figure 3. Secure state-change pipeline", "Form and AJAX actions pass method, session, CSRF, permission, business-rule and database gates.", 6.4)
                figure_added.add("secure")
            i += 1
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:].strip(), "NOTE")
            i += 1
            continue
        if re.match(r"^-\s+", line):
            active_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, re.sub(r"^-\s+", "", line))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            if active_num_id is None:
                active_num_id = create_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_num_id)
            add_rich_text(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue
        active_num_id = None
        # Q&A question lines are bold in Markdown; give them a stronger lead.
        p = doc.add_paragraph()
        if line.startswith("**") and line.endswith("**"):
            p.paragraph_format.space_before = Pt(5)
        if line.startswith(("**Showcase:**", "**Best proof:**", "**Core concept:**")):
            p.paragraph_format.space_after = Pt(4)
        add_rich_text(p, line)
        i += 1

    # Insert the team integration diagram before the final integration material.
    target = None
    for p in doc.paragraphs:
        if p.text.startswith("7. Cross-feature integration"):
            target = p
            break
    if target is not None:
        temp = doc.add_paragraph()
        add_figure(doc, "team-integration.png", "Figure 4. Three feature owners, one integrated workflow", "Tarannum Diha's club authority produces managed events, Rifat Mahmud verifies attendance and certificates, and Faisal Mahbub publishes recipient-specific engagement updates.", 6.35)
        # Move the two appended figure paragraphs to immediately after the heading.
        body = doc._element.body
        figure_nodes = list(body)[-3:]
        anchor = target._p
        for node in reversed(figure_nodes):
            anchor.addnext(node)

    # Core properties and final save.
    doc.core_properties.title = "CampusHub Team Feature Documentation"
    doc.core_properties.subject = "Architecture, ownership, demonstrations, and viva preparation"
    doc.core_properties.author = "CampusHub Project Team"
    doc.core_properties.keywords = "CampusHub, PHP, MySQL, XAMPP, database, documentation"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
